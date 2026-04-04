#!/usr/bin/env python3
"""
translate_doc.py — 文件翻譯替代工具
支援格式：.docx / .pages / .pdf（文字型 & 掃描型）
用法：python3 translate_doc.py /path/to/file --lang EN [--confirm]
"""

import argparse
import os
import shutil
import sys
import tempfile
import zipfile

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, os.path.join(BASE_DIR, 'scripts'))

from dotenv import load_dotenv
load_dotenv(os.path.join(BASE_DIR, '.env'))

from google import genai
from google.genai import types
from modules.config import GEMINI_API_KEY

client = genai.Client(api_key=GEMINI_API_KEY, http_options={'api_version': 'v1beta'})

ROOT_FOLDER_ID = '1TD7p-vdxArtQXsU27gR7jy_wQQHOmh9b'

LANG_FOLDER = {
    'EN': '英文', 'JA': '日文', 'TH': '泰文',
    'KO': '韓文', 'FR': '法文', 'DE': '德文',
    'ES': '西班牙文', 'VI': '越南文',
}

LANG_FULL = {
    'EN': 'English', 'JA': 'Japanese', 'TH': 'Thai',
    'KO': 'Korean', 'FR': 'French', 'DE': 'German',
    'ES': 'Spanish', 'VI': 'Vietnamese',
}


# ── Google Drive 工具 ─────────────────────────────────────────────────────────

def get_or_create_subfolder(svc, name, parent_id):
    q = (f"name='{name}' and '{parent_id}' in parents "
         f"and mimeType='application/vnd.google-apps.folder' and trashed=false")
    res = svc.files().list(q=q, fields='files(id)').execute()
    if res['files']:
        return res['files'][0]['id']
    meta = {
        'name': name,
        'mimeType': 'application/vnd.google-apps.folder',
        'parents': [parent_id],
    }
    return svc.files().create(body=meta, fields='id').execute()['id']


def upload_to_drive(local_path, lang_code, lang_name):
    from modules.gdrive_utils import _get_user_drive_service, upload_file
    svc = _get_user_drive_service()
    if not svc:
        print('⚠️ OAuth token 不存在，使用 service account fallback')
        folder_id = ROOT_FOLDER_ID
        return upload_file(local_path, folder_id)

    folder_name = LANG_FOLDER.get(lang_code.upper(), lang_name)
    folder_id = get_or_create_subfolder(svc, folder_name, ROOT_FOLDER_ID)

    import mimetypes
    from googleapiclient.http import MediaFileUpload
    mime_type, _ = mimetypes.guess_type(local_path)
    mime_type = mime_type or 'application/octet-stream'
    filename = os.path.basename(local_path)
    meta = {'name': filename, 'parents': [folder_id]}
    media = MediaFileUpload(local_path, mimetype=mime_type)
    f = svc.files().create(body=meta, media_body=media, fields='id, webViewLink').execute()
    print(f"⬆️  Drive 上傳完成：{folder_name}/{filename}")
    print(f"🔗 連結：{f.get('webViewLink', '（請至 Drive 查看）')}")
    return f.get('id')


# ── 前處理規則 ───────────────────────────────────────────────────────────────

import re

# 民國曆 → 西曆轉換（支援：中華民國XX年、民國XX年）
_ROC_MONTH_EN = {
    '一': 'January', '二': 'February', '三': 'March', '四': 'April',
    '五': 'May', '六': 'June', '七': 'July', '八': 'August',
    '九': 'September', '十': 'October', '十一': 'November', '十二': 'December',
}
_DIGITS = {'零':0,'一':1,'二':2,'三':3,'四':4,'五':5,'六':6,'七':7,'八':8,'九':9}
_CLASSICAL = {'十':10,'百':100}

def _roc_num_to_int(s):
    """支援逐位數字（一一五=115）和傳統中文數字（一百一十五=115）"""
    if not s:
        return None
    # 若包含十/百，使用傳統算法
    if any(c in s for c in ('十','百')):
        total, cur = 0, 0
        for ch in s:
            if ch in _CLASSICAL:
                v = _CLASSICAL[ch]
                cur = (cur or 1) * v
                if v == 100:
                    total += cur; cur = 0
            elif ch in _DIGITS:
                total += cur; cur = _DIGITS[ch]
            else:
                return None
        return total + cur
    else:
        # 逐位數字：一一五 → 1*100 + 1*10 + 5 = 115
        result = 0
        for ch in s:
            d = _DIGITS.get(ch, -1)
            if d == -1:
                return None
            result = result * 10 + d
        return result

def convert_roc_date(text):
    """偵測並轉換民國日期為西曆，不動其餘文字"""
    pattern = re.compile(
        r'(?:中華民國|民國)\s*([零一二三四五六七八九]+)\s*年'
        r'(?:\s*([一二三四五六七八九十]+(?:十[一二]?)?)\s*月)?'
        r'(?:\s*([一二三四五六七八九十]*)\s*日)?'
    )
    def replace(m):
        yr_zh = m.group(1)
        mo_zh = m.group(2)
        day_zh = m.group(3)
        yr = _roc_num_to_int(yr_zh)
        if yr is None:
            return m.group(0)
        ad_yr = yr + 1911
        month_en = _ROC_MONTH_EN.get(mo_zh, '') if mo_zh else ''
        day_int = _roc_num_to_int(day_zh) if day_zh else None
        day = str(day_int) if day_int else ''
        if month_en and day:
            return f'{month_en} {day}, {ad_yr}'
        elif month_en:
            return f'{month_en}  , {ad_yr}'
        else:
            return str(ad_yr)
    return pattern.sub(replace, text)


def pre_scan_known_translations(doc, source_detector):
    """掃描文件，找出已存在的雙語對照對（相鄰段落或同儲存格）。
    回傳：
    - pairs: {source_text: existing_translation} — 翻譯提示用
    - clear_set: set of existing_translation — 這些段落需被清空
    """
    pairs = {}
    _exclude = re.compile(r'^(TEL|FAX|E-mail|\d|http|\+)', re.IGNORECASE)

    def scan_paragraphs(paragraphs):
        texts = [p.text.strip() for p in paragraphs]
        for i, t in enumerate(texts):
            if not t or not source_detector(t):
                continue
            for j in range(i + 1, min(i + 3, len(texts))):
                candidate = texts[j]
                if (candidate
                        and not source_detector(candidate)
                        and len(candidate) > 5
                        and not _exclude.match(candidate)):
                    pairs[t] = candidate
                    break

    scan_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                scan_paragraphs(cell.paragraphs)

    clear_set = set(pairs.values())
    return pairs, clear_set


def post_process(text):
    """後處理：全形冒號 → 半形"""
    return text.replace('：', ':')


# ── Gemini 翻譯 ───────────────────────────────────────────────────────────────

def translate_texts(texts, lang_code, known=None):
    """批次翻譯文字列表，回傳等長的翻譯列表"""
    if not texts:
        return []

    lang_name = LANG_FULL.get(lang_code.upper(), lang_code)

    # 已知對照表提示
    known_hint = ''
    if known:
        known_lines = '\n'.join(f'  "{k}" → "{v}"' for k, v in known.items())
        known_hint = (
            f'\nKnown translations already in the document (use these exactly, do NOT re-translate):\n'
            f'{known_lines}\n'
        )

    numbered = '\n'.join(f'[{i}] {t}' for i, t in enumerate(texts))
    prompt = (
        f'Translate each numbered line below into {lang_name}.\n'
        f'Reply ONLY with the same numbered format [N] translation.\n'
        f'\nRules:\n'
        f'1. If an English equivalent already exists in the document, use it exactly (see known translations).\n'
        f'2. Taiwan addresses: convert to Western format → '
        f'"[Floor]F-[Unit], No. [number], [Road] Rd., [District] Dist., [City] [PostalCode], Taiwan" '
        f'(e.g. 5樓之2 → 5F-2, not "Floor 5-2")\n'
        f'3. Do NOT translate text that is already in {lang_name} — keep it as-is.\n'
        f'4. Preserve numbers, punctuation, email addresses, phone numbers.\n'
        f'5. Use half-width colon `:` not full-width `：`.\n'
        f'6. Chinese salutation 致 translates to "To:" (with colon).\n'
        f'7. Romanize road names accurately: 宜安路 = "Yi\'an Rd." (with apostrophe).\n'
        f'{known_hint}\n'
        f'{numbered}'
    )
    resp = client.models.generate_content(
        model='gemini-2.5-flash',
        contents=prompt,
    )
    lines = resp.text.strip().split('\n')
    result = [''] * len(texts)
    for line in lines:
        line = line.strip()
        if line.startswith('[') and ']' in line:
            idx_end = line.index(']')
            try:
                idx = int(line[1:idx_end])
                if 0 <= idx < len(result):
                    content = line[idx_end + 1:].lstrip('] ').strip()
                    result[idx] = post_process(content)
            except ValueError:
                pass
    for i, t in enumerate(result):
        if not t:
            result[i] = texts[i]
    return result


def has_chinese(text):
    return any('\u4e00' <= c <= '\u9fff' for c in text)

def has_thai(text):
    return any('\u0e00' <= c <= '\u0e7f' for c in text)

def make_source_detector(doc):
    """從文件內容自動偵測主要來源語言，回傳 (detector_fn, lang_name)"""
    all_text = ' '.join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    all_text += ' ' + p.text
    if has_chinese(all_text):
        return has_chinese, 'Chinese'
    if has_thai(all_text):
        return has_thai, 'Thai'
    # 預設：英文（含拉丁字母）
    return lambda t: bool(re.search(r'[A-Za-z]', t)), 'English'


# ── DOCX 處理 ─────────────────────────────────────────────────────────────────

def translate_docx(src_path, lang_code, out_path):
    from docx import Document

    doc = Document(src_path)

    # 自動偵測來源語言
    source_detector, src_lang_name = make_source_detector(doc)
    print(f'🔍 偵測來源語言：{src_lang_name}')

    # 預掃描文件內已存在的雙語對照
    known, clear_set = pre_scan_known_translations(doc, source_detector)
    if known:
        print(f'📚 找到 {len(known)} 組已知對照（來源位置翻譯，對應位置清除）：{list(known.keys())}')

    # 用 XPath 路徑字串去重：比 id() 更穩定，不受 lxml proxy GC/重用影響
    _tree = doc.element.getroottree()
    processed_paths = set()

    def process_paragraphs(paragraphs):
        for p in paragraphs:
            path = _tree.getpath(p._p)
            if path in processed_paths:
                continue
            processed_paths.add(path)

            full_text = p.text
            full_stripped = full_text.strip()
            if not full_stripped:
                continue

            # 已存在的翻譯對照 → 清空（避免重複）
            if full_stripped in clear_set:
                for run in p.runs:
                    run.text = ''
                continue

            # 非來源語言 → 不動（已是目標語言或中性格式）
            if not source_detector(full_text):
                continue

            # 民國日期：段落層級偵測後整段轉換
            converted = convert_roc_date(full_text)
            if converted != full_text:
                if p.runs:
                    p.runs[0].text = converted
                    for run in p.runs[1:]:
                        run.text = ''
                continue

            # 來源語言文字：整段翻譯（Gemini 保留完整上下文）
            translated_para = translate_texts([full_stripped], lang_code, known=known)
            if translated_para:
                if p.runs:
                    p.runs[0].text = translated_para[0]
                    for run in p.runs[1:]:
                        run.text = ''
                else:
                    # 段落無直接 <w:r> run（文字在 hyperlink/field 元素中）
                    # → 清除非 pPr 子元素，加入新 run
                    from docx.oxml.ns import qn
                    pPr = p._p.find(qn('w:pPr'))
                    for child in list(p._p):
                        p._p.remove(child)
                    if pPr is not None:
                        p._p.append(pPr)
                    p.add_run(translated_para[0])

        # 格式後處理：全形冒號統一換成半形（純格式，非翻譯）
        for p in paragraphs:
            if '：' in p.text:
                for run in p.runs:
                    if '：' in run.text:
                        run.text = run.text.replace('：', ':')

    process_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs)

    doc.save(out_path)
    print(f'✅ 已儲存：{out_path}')


# ── PAGES 處理 ────────────────────────────────────────────────────────────────

def translate_pages(src_path, lang_code, out_path):
    """解壓 .pages（zip）→ 修改 index.xml → 重新打包"""
    import re

    with tempfile.TemporaryDirectory() as tmpdir:
        with zipfile.ZipFile(src_path) as z:
            z.extractall(tmpdir)

        index_path = os.path.join(tmpdir, 'index.xml')
        if not os.path.exists(index_path):
            # Pages '13+ 用 Index/Document.iwa (protobuf)，無法直接解析
            print('⚠️ 此 .pages 格式使用 protobuf，暫不支援。請用 Pages App 另存為 .docx 後再試。')
            return None

        with open(index_path, 'r', encoding='utf-8') as f:
            xml = f.read()

        # 找出所有文字節點
        pattern = re.compile(r'(<sf:s[^>]*>)(.*?)(</sf:s>)', re.DOTALL)
        matches = list(pattern.finditer(xml))
        chinese_matches = [(m, m.group(2)) for m in matches if has_chinese(m.group(2))]

        if not chinese_matches:
            print('ℹ️  未發現中文文字。')
            shutil.copy2(src_path, out_path)
            return out_path

        print(f'🔄 翻譯 {len(chinese_matches)} 個文字段...')
        texts = [t for _, t in chinese_matches]
        translated = translate_texts(texts, lang_code)

        # 倒序替換（避免 offset 偏移）
        for (m, _), new_text in reversed(list(zip(chinese_matches, translated))):
            xml = xml[:m.start(2)] + new_text + xml[m.end(2):]

        with open(index_path, 'w', encoding='utf-8') as f:
            f.write(xml)

        with zipfile.ZipFile(out_path, 'w', zipfile.ZIP_DEFLATED) as zout:
            for root, _, files in os.walk(tmpdir):
                for fname in files:
                    fpath = os.path.join(root, fname)
                    zout.write(fpath, os.path.relpath(fpath, tmpdir))

    print(f'✅ 已儲存：{out_path}')
    return out_path


# ── PDF 處理 ──────────────────────────────────────────────────────────────────

def is_scanned_pdf(pdf_path):
    """若每頁文字量極少，判定為掃描型"""
    import fitz
    doc = fitz.open(pdf_path)
    total_chars = sum(len(page.get_text()) for page in doc)
    doc.close()
    return total_chars < 50 * len(fitz.open(pdf_path))


def translate_text_pdf(src_path, lang_code, out_path):
    """文字型 PDF：提取文字 → 翻譯 → reportlab 重建"""
    import fitz
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import pt
    from reportlab.pdfgen import canvas
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    # 嘗試載入中文字體（macOS 內建）
    font_path = '/System/Library/Fonts/PingFang.ttc'
    font_name = 'PingFang'
    if os.path.exists(font_path):
        pdfmetrics.registerFont(TTFont(font_name, font_path))
    else:
        font_name = 'Helvetica'

    doc = fitz.open(src_path)
    c = canvas.Canvas(out_path)

    for page in doc:
        w, h = page.rect.width, page.rect.height
        c.setPageSize((w, h))

        blocks = page.get_text('blocks')
        texts = [b[4].strip() for b in blocks if b[4].strip() and has_chinese(b[4])]

        if texts:
            print(f'🔄 翻譯第 {page.number + 1} 頁 {len(texts)} 個段落...')
            translated = translate_texts(texts, lang_code)
            t_map = dict(zip(texts, translated))
        else:
            t_map = {}

        for b in blocks:
            x0, y0, x1, y1, text, *_ = b
            if not text.strip():
                continue
            display = t_map.get(text.strip(), text.strip())
            c.setFont(font_name, 10)
            # PDF 座標系 y 軸反轉
            c.drawString(x0 * pt, (h - y1) * pt, display)

        c.showPage()

    doc.close()
    c.save()
    print(f'✅ 已儲存：{out_path}')


def translate_scanned_pdf(src_path, lang_code, out_path, confirm_mode=False):
    """掃描型 PDF：截圖背景 + Gemini 翻譯文字疊加"""
    import fitz
    from reportlab.lib.units import pt
    from reportlab.pdfgen import canvas
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    font_path = '/System/Library/Fonts/PingFang.ttc'
    font_name = 'PingFang'
    if os.path.exists(font_path):
        pdfmetrics.registerFont(TTFont(font_name, font_path))
    else:
        font_name = 'Helvetica'

    doc = fitz.open(src_path)
    c = canvas.Canvas(out_path)

    for page in doc:
        w, h = page.rect.width, page.rect.height
        c.setPageSize((w, h))

        # 渲染頁面截圖作為背景
        pix = page.get_pixmap(dpi=150)
        img_path = f'/tmp/_scan_page_{page.number}.png'
        pix.save(img_path)
        c.drawImage(img_path, 0, 0, width=w, height=h)

        if not confirm_mode:
            # 用 Gemini Vision 提取文字並翻譯
            lang_name = LANG_FULL.get(lang_code.upper(), lang_code)
            img_part = types.Part.from_bytes(data=open(img_path, 'rb').read(), mime_type='image/png')
            prompt = (
                f'Extract all text from this scanned document page and translate it to {lang_name}. '
                f'Reply in this format for each text block:\n'
                f'[x%,y%] translated text\n'
                f'Where x% and y% are the approximate position as percentage of page width/height (top-left origin). '
                f'Keep numbers, proper nouns, and non-source-language text as-is.'
            )
            resp = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=[prompt, img_part],
            )

            import re
            for line in resp.text.strip().split('\n'):
                m = re.match(r'\[(\d+)%,(\d+)%\]\s+(.*)', line.strip())
                if m:
                    x_pct, y_pct, text = float(m.group(1)), float(m.group(2)), m.group(3)
                    x = w * x_pct / 100
                    y = h * (1 - y_pct / 100)
                    c.setFont(font_name, 9)
                    c.setFillColorRGB(0, 0, 0.6, alpha=0.85)
                    c.drawString(x, y, text)

        os.remove(img_path)
        c.showPage()

    doc.close()
    c.save()
    print(f'✅ 已儲存：{out_path}')


# ── 主流程 ────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument('file', help='來源文件路徑')
    parser.add_argument('--lang', required=True, help='目標語言代碼（EN/JA/TH/KO/FR/DE/ES/VI）')
    parser.add_argument('--confirm', action='store_true',
                        help='掃描PDF確認模式：只產生截圖預覽，不疊加翻譯文字')
    parser.add_argument('--upload', action='store_true', default=False,
                        help='完成後上傳至 Google Drive')
    args = parser.parse_args()

    src = os.path.expanduser(args.file)
    if not os.path.exists(src):
        print(f'❌ 找不到檔案：{src}')
        sys.exit(1)

    lang = args.lang.upper()
    lang_name = LANG_FOLDER.get(lang, lang)
    ext = os.path.splitext(src)[1].lower()
    base = os.path.splitext(os.path.basename(src))[0]
    suffix = '_preview' if args.confirm else f'_{lang}'
    out_path = f'/tmp/{base}{suffix}{ext}'

    print(f'📄 來源：{src}')
    print(f'🌐 目標語言：{lang_name}（{lang}）')

    if ext == '.docx':
        translate_docx(src, lang, out_path)

    elif ext == '.pages':
        result = translate_pages(src, lang, out_path)
        if result is None:
            sys.exit(1)

    elif ext == '.pdf':
        if is_scanned_pdf(src):
            print('📷 偵測到掃描型 PDF')
            translate_scanned_pdf(src, lang, out_path, confirm_mode=args.confirm)
            if args.confirm:
                print(f'\n📋 預覽檔已產生：{out_path}')
                print('確認排版後，請重新執行並加上 --upload 參數上傳。')
                return
        else:
            translate_text_pdf(src, lang, out_path)
    else:
        print(f'❌ 不支援的格式：{ext}')
        sys.exit(1)

    if args.upload:
        upload_to_drive(out_path, lang, lang_name)
    else:
        print(f'\n📁 輸出檔案：{out_path}')
        print('加上 --upload 參數可直接上傳至 Google Drive。')


if __name__ == '__main__':
    main()
