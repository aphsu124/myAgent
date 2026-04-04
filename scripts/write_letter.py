#!/usr/bin/env python3
"""
write_letter.py — 在信紙抬頭下方寫入信件內文
用法：
  python3 scripts/write_letter.py \
    --letterhead /path/to/信紙.docx \
    --content /tmp/letter_content.json \
    --out /tmp/output.docx \
    [--upload] [--folder 英文]

content JSON 格式（段落列表）：
  [
    {"text": "To Thailand Trade and Economic Office", "align": "left"},
    {"text": "", "align": "left"},
    {"text": "13/April/2026", "align": "center"},
    {"text": "\\tThe company...", "align": "justify"},
    ...
  ]
  align 值：left / center / right / justify
  可選欄位：size（字級，預設 12）、bold（布林，預設 false）
"""

import argparse
import json
import os
import sys

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, os.path.join(BASE_DIR, 'scripts'))

from dotenv import load_dotenv
load_dotenv(os.path.join(BASE_DIR, '.env'))

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

ROOT_FOLDER_ID = '1TD7p-vdxArtQXsU27gR7jy_wQQHOmh9b'

ALIGN_MAP = {
    'left':    WD_ALIGN_PARAGRAPH.LEFT,
    'center':  WD_ALIGN_PARAGRAPH.CENTER,
    'right':   WD_ALIGN_PARAGRAPH.RIGHT,
    'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
}

HEADER_BLANK_LINES = 4  # 信紙表格下方固定空行數


def _clear_paragraph(p):
    """清除段落內容，保留段落格式（pPr）"""
    pPr = p._p.find(qn('w:pPr'))
    for child in list(p._p):
        p._p.remove(child)
    if pPr is not None:
        p._p.append(pPr)


def _set_paragraph(p, text, align, size=12, bold=False):
    """將段落設為指定文字、對齊、字級"""
    _clear_paragraph(p)
    if text:
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
    p.alignment = ALIGN_MAP.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


def write_on_letterhead(letterhead_path, content_items, out_path):
    """
    在信紙表格後方插入內文。
    content_items: list of dict，每個 dict 代表一個段落。
    """
    doc = Document(letterhead_path)

    # 找最後一個表格在文件 body XML 中的位置
    body = doc.element.body
    last_tbl = None
    for child in body:
        if child.tag == qn('w:tbl'):
            last_tbl = child

    if last_tbl is None:
        print('⚠️  找不到表格，從文件開頭寫入')

    # 收集表格後的所有段落
    body_children = list(body)
    if last_tbl is not None:
        tbl_idx = body_children.index(last_tbl)
        post_tbl_paras = [
            c for c in body_children[tbl_idx + 1:]
            if c.tag == qn('w:p')
        ]
    else:
        post_tbl_paras = [c for c in body_children if c.tag == qn('w:p')]

    # 建立 Paragraph 包裝物件
    from docx.text.paragraph import Paragraph
    post_paras = [Paragraph(p, body) for p in post_tbl_paras]

    # 所需總段落數 = 空行 + 內文段落
    total_needed = HEADER_BLANK_LINES + len(content_items)

    # 若現有段落不足，補充
    while len(post_paras) < total_needed:
        new_p = doc.add_paragraph()
        post_paras.append(new_p)

    # 寫入空行
    for i in range(HEADER_BLANK_LINES):
        _set_paragraph(post_paras[i], '', 'left')

    # 寫入內文段落
    for i, item in enumerate(content_items):
        p = post_paras[HEADER_BLANK_LINES + i]
        _set_paragraph(
            p,
            text=item.get('text', ''),
            align=item.get('align', 'left'),
            size=item.get('size', 12),
            bold=item.get('bold', False),
        )

    # 清空剩餘多餘段落
    for p in post_paras[total_needed:]:
        _clear_paragraph(p)

    doc.save(out_path)
    print(f'✅ 已儲存：{out_path}')


def upload_to_drive(local_path, folder_name):
    from modules.gdrive_utils import _get_user_drive_service
    import mimetypes
    from googleapiclient.http import MediaFileUpload

    # inline import of get_or_create_subfolder
    def get_or_create_subfolder(svc, name, parent_id):
        q = (f"name='{name}' and '{parent_id}' in parents "
             f"and mimeType='application/vnd.google-apps.folder' and trashed=false")
        res = svc.files().list(q=q, fields='files(id)').execute()
        if res['files']:
            return res['files'][0]['id']
        meta = {
            'name': name,
            'mimeType': 'application/vnd.google-apps.folder',
            'parents': [parent_id],
        }
        return svc.files().create(body=meta, fields='id').execute()['id']

    svc = _get_user_drive_service()
    if not svc:
        print('⚠️  OAuth token 不存在，無法上傳')
        return None

    folder_id = get_or_create_subfolder(svc, folder_name, ROOT_FOLDER_ID)
    mime_type, _ = mimetypes.guess_type(local_path)
    mime_type = mime_type or 'application/octet-stream'
    filename = os.path.basename(local_path)
    meta = {'name': filename, 'parents': [folder_id]}
    media = MediaFileUpload(local_path, mimetype=mime_type)
    f = svc.files().create(body=meta, media_body=media, fields='id,webViewLink').execute()
    print(f'⬆️  Drive 上傳完成：{folder_name}/{filename}')
    print(f'🔗 連結：{f.get("webViewLink")}')
    return f.get('webViewLink')


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument('--letterhead', required=True, help='信紙 .docx 路徑')
    parser.add_argument('--content', required=True, help='內文 JSON 檔路徑')
    parser.add_argument('--out', required=True, help='輸出 .docx 路徑')
    parser.add_argument('--upload', action='store_true', help='完成後上傳至 Drive')
    parser.add_argument('--folder', default='英文', help='Drive 子資料夾名稱（預設：英文）')
    args = parser.parse_args()

    src = os.path.expanduser(args.letterhead)
    if not os.path.exists(src):
        print(f'❌ 找不到信紙檔案：{src}')
        sys.exit(1)

    with open(args.content, 'r', encoding='utf-8') as f:
        content_items = json.load(f)

    write_on_letterhead(src, content_items, args.out)

    if args.upload:
        upload_to_drive(args.out, args.folder)


if __name__ == '__main__':
    main()
